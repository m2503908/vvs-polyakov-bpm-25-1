from dialog import Ui_MainWindow

import sys
import datetime as dt
from random import randint

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog


class MyWidget(QMainWindow, Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.lineEdit_2.setText('.'.join(str(dt.date.today()).split('-')[::-1]))
        self.pushButton.clicked.connect(self.run)

    def run(self):
        user_name = self.lineEdit.text()
        user_post = self.lineEdit_3.text()
        type_of_doc = self.comboBox.currentText()
        director_name = self.comboBox_2.currentText()
        action_date = dt.datetime(self.calendarWidget.selectedDate().year(),
                                  self.calendarWidget.selectedDate().month(),
                                  self.calendarWidget.selectedDate().day())
        action_date = '.'.join(str(action_date.date()).split('-')[::-1])
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        para = doc.add_paragraph('Директору школы №1580')
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para = doc.add_paragraph(str(director_name))
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para = doc.add_paragraph(f'от {str(user_post)}')
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para = doc.add_paragraph(str(user_name))
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para = doc.add_paragraph('Заявление')
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num_of_days = self.lineEdit_4.text()
        if str(self.comboBox.currentText()) == "Оплачиваемый отпуск":
            s1 = 'Прошу предоставить мне оплачиваемый отпуск с '
            s2 = action_date
            s3 = f'г. продолжительностью {num_of_days} календарных дней.'
            para = doc.add_paragraph(s1 + s2 + s3)
        elif str(self.comboBox.currentText()) == "Отпуск за свой счет":
            s1 = 'Прошу предоставить мне отпуск без сохранения заработной платы с '
            s2 = action_date
            s3 = f'г. продолжительностью {num_of_days} календарных дня.'
            para = doc.add_paragraph(s1 + s2 + s3)
        elif str(self.comboBox.currentText()) == "Увольнение":
            s1 = 'Прошу в соответствии со статьей 80 Трудового кодекса уволить меня по '
            s2 = 'собственному желанию с занимаемой должности с '
            s3 = action_date
            para = doc.add_paragraph(s1 + s2 + s3)
            self.lineEdit_4.setReadOnly()
        else:
            s1 = 'Прошу Вас рассмотреть возможность повышения моей заработной платы поскольку '
            s2 = 'объем выполняемых работ был увеличен. С '
            s3 = action_date
            s4 = ' объем выполняемых работ был увеличен на 10%. В связи с этим прошу рассмотреть'
            s5 = 'возможность повышения заработной платы соответственно на 10%.'
            para = doc.add_paragraph(s1 + s2 + s3 + s4 + s5)
            self.lineEdit_4.setReadOnly()
        sp = '                                                                                     '
        para = doc.add_paragraph(' ')
        para = doc.add_paragraph(f'{str(self.lineEdit_2.text())}{sp}__________')
        para = doc.add_paragraph(str(user_name))
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para = doc.add_paragraph(' ')
        para = doc.add_paragraph(' ')
        para = doc.add_paragraph(f'{str(self.lineEdit_2.text())}{sp}__________')
        para = doc.add_paragraph(str(director_name))
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        file_name = f'document{randint(1, 100000)}.docx'
        doc.save(file_name)

        sys.exit(app.exec())


app = QApplication(sys.argv)
ex = MyWidget()
ex.show()
sys.exit(app.exec())
